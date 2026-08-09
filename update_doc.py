# -*- coding: utf-8 -*-
"""更新 里程碑项目思路.docx，补充测试设计思路内容"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r"D:\MyDesktop\里程碑项目思路.docx"
doc = Document(doc_path)

# 添加分页
doc.add_page_break()

# 标题
title = doc.add_heading('Adaptive LP Vault 项目测试设计思路与教学', level=1)

# ==================== 第一部分：测试文件设计思路 ====================
doc.add_heading('一、各测试文件的设计思路', level=2)

doc.add_heading('1. BaseTest.t.sol - 测试基类', level=3)
doc.add_paragraph(
    '设计思路：模拟完整的生产环境部署，为所有测试提供统一的setUp。'
    '包括部署Mock WETH/USDC、Mock Uniswap V2/V3（两个费率池）、设置初始价格2000 USDC/ETH、'
    '部署治理/策略/预言机/金库/适配器/激励合约、给三个测试用户mint代币并授权。'
    '所有测试合约继承BaseTest，避免重复部署代码。'
)

doc.add_heading('2. VaultUnitTest.t.sol - 金库核心测试', level=3)
doc.add_paragraph(
    '设计思路：覆盖AdaptiveLPVault的所有核心功能，按模块组织：'
)
items = [
    '存款测试：双币存款、单币存款、多用户存款、小额灰尘、ERC4626标准接口',
    '取款测试：withdrawDual正常/部分/全部、滑点保护、多用户隔离、零份额回滚',
    '重平衡测试：首次重平衡、多次重平衡、高波动、冷却期、激励失败容错',
    '权限测试：onlyOwner、setPaused、参数设置边界',
    '视图函数：totalAssets、getDistribution、convertToShares/Assets',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3. ComponentsTest.t.sol - 组件独立测试', level=3)
doc.add_paragraph(
    '设计思路：对每个核心组件进行独立单元测试，不依赖金库：'
)
items = [
    'V2AdapterTest：addLiquidity/removeLiquidity/collectFees/withdrawAll',
    'V3AdapterTest：多区间mint/burn/collect、position管理',
    'OracleTest：TWAP计算、价格换算、窗口设置',
    'StrategyTest：波动率分档、权重分配、tick计算',
    'IncentivesTest：激励计算、冷却期、奖励领取',
    'GovernanceTest：提案/投票/执行流程',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4. UniswapMathTest.t.sol - 数学库测试', level=3)
doc.add_paragraph(
    '设计思路：对UniswapMath库的纯函数进行fuzz测试，'
    '包括sqrt计算、tick和sqrtPrice转换、价格换算等，'
    '用不同输入验证数学正确性。'
)

doc.add_heading('5. IntegrationTest.t.sol - 集成测试', level=3)
doc.add_paragraph(
    '设计思路：端到端测试完整用户旅程：'
    'Alice存款→Bob存款→执行重平衡→价格变动→再次重平衡→Alice部分赎回→Bob全部赎回。'
    '验证多用户场景下的份额计算、资产隔离、滑点保护是否正确。'
)

doc.add_heading('6. ForkTest.t.sol - 主网分叉测试', level=3)
doc.add_paragraph(
    '设计思路：在Sepolia/主网分叉上用真实Uniswap合约测试，'
    '验证适配器与真实合约的兼容性，发现Mock覆盖不到的问题。'
    '注意：需要RPC节点，默认不运行。'
)

# ==================== 第二部分：如何编写Solidity测试 ====================
doc.add_heading('二、如何为本项目编写Solidity测试合约', level=2)

doc.add_heading('1. setUp怎么设计', level=3)
doc.add_paragraph(
    'setUp是每个测试合约运行前的准备函数，等价于JUnit的@BeforeEach。'
    '本项目推荐继承BaseTest，它已经完成了完整部署：'
)
code1 = '''contract MyTest is BaseTest {
    function setUp() public override {
        super.setUp();  // 调用基类setUp，部署所有合约
        // 这里可以添加自己的额外初始化
    }
}'''
doc.add_paragraph(code1, style='No Spacing')

doc.add_paragraph('setUp设计原则：')
items = [
    '每个测试函数运行前都会重新调用setUp，状态完全隔离',
    '不要在setUp里执行太多业务逻辑，保持测试独立性',
    '使用makeAddr("name")生成可读的测试地址',
    '初始代币数量要足够大，避免测试中余额不足',
    '授权额度设为type(uint256).max，避免重复approve',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2. 正常用例怎么写', level=3)
doc.add_paragraph('标准模式：Arrange-Act-Assert')
code2 = '''function test_Deposit_Success() public {
    // Arrange: 准备
    uint256 wethAmt = 1 ether;
    uint256 usdcAmt = 2000e6;
    uint256 balBefore = vault.balanceOf(alice);

    // Act: 执行
    vm.startPrank(alice);
    uint256 shares = vault.deposit(wethAmt, usdcAmt, 0);
    vm.stopPrank();

    // Assert: 验证
    assertGt(shares, 0, "should mint shares");
    assertEq(vault.balanceOf(alice), balBefore + shares);
    assertEq(vault.totalSupply(), shares);
}'''
doc.add_paragraph(code2, style='No Spacing')

doc.add_heading('3. 异常回滚用例怎么写', level=3)
doc.add_paragraph('使用vm.expectRevert()验证预期回滚：')
code3 = '''function test_Revert_ZeroAmount() public {
    vm.startPrank(alice);
    // 不指定错误信息，只要回滚就通过
    vm.expectRevert();
    vault.deposit(0, 0, 0);
    vm.stopPrank();
}

function test_Revert_SlippageTooHigh() public {
    uint256 shares = _deposit(alice, 10 ether, 20000e6);
    vm.startPrank(alice);
    // 指定具体的custom error selector
    vm.expectRevert(AdaptiveLPVault.SlippageExceeded.selector);
    vault.withdrawDual(shares, 100 ether, 200000e6);
    vm.stopPrank();
}'''
doc.add_paragraph(code3, style='No Spacing')

doc.add_heading('4. 滑点类场景如何测试', level=3)
doc.add_paragraph(
    '滑点测试的关键是：不要传min=0！要根据合约状态真实计算预期输出：'
)
code4 = '''function test_Withdraw_RealisticSlippage() public {
    // 1. 先存款
    uint256 shares = _deposit(alice, 10 ether, 20000e6);

    // 2. 执行重平衡让资金进入adapter
    vm.prank(alice);
    vault.rebalance();

    // 3. 根据getDistribution计算预期输出
    (uint256 iW, uint256 iU, uint256 v2W, uint256 v2U,
     uint256 v3LW, uint256 v3LU, uint256 v3HW, uint256 v3HU) = vault.getDistribution();
    uint256 totalWeth = iW + v2W + v3LW + v3HW;
    uint256 totalUsdc = iU + v2U + v3LU + v3HU;
    uint256 totalSup = vault.totalSupply();

    // 4. 按份额比例计算预期，设置99%作为min（1%滑点容忍）
    uint256 expWeth = totalWeth * shares / totalSup;
    uint256 expUsdc = totalUsdc * shares / totalSup;
    uint256 minWeth = expWeth * 99 / 100;
    uint256 minUsdc = expUsdc * 99 / 100;

    // 5. 记录赎回前余额
    uint256 wethBefore = weth.balanceOf(alice);
    uint256 usdcBefore = usdc.balanceOf(alice);

    // 6. 执行赎回
    vm.startPrank(alice);
    (uint256 outW, uint256 outU) = vault.withdrawDual(shares, minWeth, minUsdc);
    vm.stopPrank();

    // 7. 验证输出在合理范围内
    assertGe(outW, minWeth);
    assertLe(outW, expWeth);  // 不超过理论值（可能有少量手续费）
    assertGe(outU, minUsdc);
}'''
doc.add_paragraph(code4, style='No Spacing')

doc.add_heading('5. 常用Cheatcodes速查', level=3)
cheats = [
    ('vm.prank(addr)', '下一次调用以addr身份执行'),
    ('vm.startPrank(addr) / vm.stopPrank()', '一段代码以addr身份执行'),
    ('vm.warp(timestamp)', '设置block.timestamp'),
    ('vm.roll(blockNumber)', '设置block.number'),
    ('skip(seconds)', '快进时间（等价于warp）'),
    ('vm.expectRevert()', '下一次调用预期回滚'),
    ('vm.expectRevert(Error.selector)', '预期回滚且错误匹配'),
    ('vm.assume(condition)', 'fuzz测试过滤无效输入'),
    ('makeAddr("name")', '生成确定性命名地址'),
    ('assertEq(a, b)', '断言相等'),
    ('assertGt/Ge/Lt/Le', '断言大小关系'),
    ('assertApproxEqRel(a, b, tol)', '断言近似相等（相对误差）'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Cheatcode'
hdr[1].text = '用途'
for code, desc in cheats:
    row = table.add_row().cells
    row[0].text = code
    row[1].text = desc

# ==================== 第三部分：注意事项与踩坑 ====================
doc.add_heading('三、本项目测试注意事项与常见踩坑', level=2)

doc.add_heading('坑1：withdraw传min=0导致测试漏洞', level=3)
doc.add_paragraph(
    '问题：早期测试中withdrawDual的minWETH/minUSDC参数全传0，'
    '虽然测试都通过，但没有真正覆盖滑点校验逻辑。'
    '线上用户实际操作时，合约的滑点保护会触发revert。'
)
doc.add_paragraph(
    '正确做法：根据getDistribution()和totalSupply()计算预期输出，'
    '设置合理的min（如95-99%）。同时添加反向测试：min设过高时应revert。'
)

doc.add_heading('坑2：token0/token1顺序问题', level=3)
doc.add_paragraph(
    '问题：Uniswap V3的price = token1/token0，token顺序由地址大小决定。'
    'Sepolia上USDC地址(0x1c7D...) < WETH地址(0xfFf9...)，所以token0=USDC。'
    '但本地Mock的createPool不排序，BaseTest中createPool(weth, usdc)，所以token0=WETH。'
    '如果测试代码硬编码token0=USDC，本地和Sepolia行为不一致。'
)
doc.add_paragraph(
    '正确做法：合约中通过TOKEN0_IS_WETH immutable动态判断；'
    '测试中不要假设token顺序，从pool.token0()/token1()读取。'
)

doc.add_heading('坑3：BigNumber.toNumber()大数溢出', level=3)
doc.add_paragraph(
    '问题：JavaScript的Number安全整数上限是2^53 ≈ 9×10^15。'
    'sqrtPriceX96约为3.5×10^24，平方后更大，toNumber()会丢失精度导致价格计算错误。'
)
doc.add_paragraph(
    '正确做法：用ethers.BigNumber做所有运算，最后转字符串手动处理小数位。'
)

doc.add_heading('坑4：Mock与真实合约行为差异', level=3)
doc.add_paragraph(
    '问题：Mock合约为简化测试，行为与真实合约有差异：'
)
items = [
    'MockUniswapV3Factory.createPool不排序token（真实V3 Factory会排序）',
    'MockV3Pool.observe假设价格恒定（真实observe返回历史tick累积值）',
    'MockV2 addLiquidity直接用传入值（真实Router会计算最优LP数量）',
    'MockV3手续费是硬编码的（真实手续费来自交易）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph(
    '正确做法：Mock用于快速单元测试没问题，但关键逻辑要在主网分叉上验证。'
)

doc.add_heading('坑5：重平衡冷却期', level=3)
doc.add_paragraph(
    '问题：两次rebalance之间有冷却期（默认300秒）。'
    '测试中连续调用rebalance会被revert。'
    '第一次rebalance特殊处理，不需要冷却。'
)
doc.add_paragraph(
    '正确做法：测试中用skip(301)或vm.warp快进时间。'
)

doc.add_heading('坑6：份额小数位是6位不是18位', level=3)
doc.add_paragraph(
    '问题：ALP份额代币继承ERC4626，以USDC为asset，decimals=6。'
    '如果测试中按18位精度计算份额会出错。'
)
doc.add_paragraph(
    '正确做法：parseUnits("100", 6)表示100份额，余额断言时注意单位。'
)

doc.add_heading('坑7：Test文件中的辅助函数', level=3)
doc.add_paragraph(
    'BaseTest提供了辅助函数：_deposit(addr, weth, usdc)内部处理prank和授权，'
    '返回shares。测试中优先使用这些辅助函数保持代码简洁。'
)

# ==================== 第四部分：避免"测试全过但线上报错" ====================
doc.add_heading('四、如何避免"测试全过但线上报错"', level=2)

items = [
    ('不要为了通过测试而传边界值',
     'min=0、amount=type(uint256).max这类值虽然能让测试通过，但不代表真实用户场景。'
     '要模拟真实用户行为，计算合理的参数。'),
    ('测试要覆盖异常路径',
     '不仅测试"正常工作"，更要测试"出错时是否正确回滚"。'
     '每个require/error都应该有对应的expectRevert测试。'),
    ('多用户场景必须测',
     '单用户存款-取款通过不代表多用户无问题。'
     '要测试A存款B存款A取款后B的资产是否正确，验证用户间无交叉污染。'),
    ('状态变化后全面验证',
     'rebalance后、withdraw后不要只测返回值，'
     '还要验证totalAssets、totalSupply、getDistribution、其他用户余额。'),
    ('注意Mock覆盖不到的地方',
     'Mock简化了很多逻辑，真实合约的回调、手续费累积、价格更新等Mock没有。'
     '关键路径要写fork测试用真实合约验证。'),
    ('事件验证',
     '重要状态变化应该触发事件，测试中用vm.expectEmit验证事件参数正确。'
     '这能捕获"状态改了但事件没发"或"事件参数错误"的问题。'),
    ('Fuzz测试发现边界case',
     '对于数学计算、权重分配等函数，用fuzz测试大量随机输入，'
     '能发现手工测试想不到的边界条件。'),
]
for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(title_text + '：')
    run.bold = True
    p.add_run(desc)

# 保存
doc.save(doc_path)
print("文档更新成功！")
